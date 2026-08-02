"""
DSFD - Générateur de Rapport v2 — Structure exacte du canevas
"""
import io
from docx import Document
from docx.shared import RGBColor

NOIR = RGBColor(0x00, 0x00, 0x00)

def _ecrire(cell, texte):
    """Écrit dans une cellule en effaçant le contenu rouge."""
    if not cell.paragraphs:
        return
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(str(texte or ''))
    run.font.color.rgb = NOIR

def _remplacer_dans_doc(doc, replacements):
    for para in doc.paragraphs:
        for run in para.runs:
            for old, new in replacements.items():
                if old in run.text:
                    run.text = run.text.replace(old, str(new))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        for old, new in replacements.items():
                            if old in run.text:
                                run.text = run.text.replace(old, str(new))


# ══ TABLE 2 — Infos générales SFD (index 1) ══
def _table2(doc, mission):
    t = doc.tables[1]
    infos = mission.get('infos_sfd', {})
    mapping = {
        'Date de constitution':       infos.get('date_constitution', ''),
        'immatriculation':            infos.get('date_rccm', ''),
        'Numéro RCCM':                infos.get('num_rccm', ''),
        "Date d'agrément":            infos.get('date_agrement', ''),
        'agrément':                   infos.get('num_agrement', ''),
        'Situation géographique':     infos.get('situation_geo', ''),
        'démarrage':                  infos.get('date_demarrage', ''),
        'Adresse':                    infos.get('adresse', ''),
        'Contacts':                   infos.get('contacts', ''),
        'points de service':          infos.get('nb_points', ''),
        'Activités du SFD':           infos.get('activites', ''),
    }
    for row in t.rows:
        if len(row.cells) < 2:
            continue
        label = row.cells[0].text.strip()
        for key, val in mapping.items():
            if key.lower() in label.lower() and val:
                _ecrire(row.cells[1], val)
                break


# ══ TABLE 3 — Membres CA (index 2) ══
def _table3(doc, mission):
    t = doc.tables[2]
    membres = mission.get('organes', {}).get('ca', [])
    # Lignes de données commencent à L2 (L0=header, L1=titre CA)
    data_rows = [r for r in t.rows if r.cells[0].text.strip().isdigit()]
    for i, m in enumerate(membres):
        if i >= len(data_rows): break
        cells = data_rows[i].cells
        if len(cells) >= 5:
            _ecrire(cells[1], m.get('nom', ''))
            _ecrire(cells[2], m.get('fonction', ''))
            _ecrire(cells[3], m.get('debut', ''))
            _ecrire(cells[4], m.get('contact', ''))


# ══ TABLE 4 — Personnel (index 3) ══
def _table4(doc, mission):
    t = doc.tables[3]
    personnel = mission.get('personnel', [])
    data_rows = [r for r in t.rows if r.cells[0].text.strip().isdigit()]
    for i, p in enumerate(personnel):
        if i >= len(data_rows): break
        cells = data_rows[i].cells
        if len(cells) >= 5:
            _ecrire(cells[1], p.get('nom', ''))
            _ecrire(cells[2], p.get('fonction', ''))
            _ecrire(cells[3], p.get('date_embauche', ''))
            _ecrire(cells[4], p.get('contrat', ''))


# ══ TABLE 5 — Indicateurs financiers (index 4) ══
def _table5(doc, mission):
    t = doc.tables[4]
    ind = mission.get('indicateurs_financiers', {})
    if not ind: return

    periodes = ind.get('periodes', ['', '', ''])
    # Ligne 0 = header avec périodes (colonnes 1,2,3)
    if len(t.rows) > 1:
        header = t.rows[1]
        for j, p in enumerate(periodes[:3]):
            if j+1 < len(header.cells):
                _ecrire(header.cells[j+1], p)

    label_map = {
        'Nombre de membres':          'nb_membres',
        'Trésorerie':                 'tresorerie',
        'Capital social':             'capital_social',
        'dépôts':                     'encours_depots',
        'Dépôt moyen':                'depot_moyen',
        'crédits octroyés':           'nb_credits',
        'Production de prêts':        'production_prets',
        "Charges d'exploitation":     'charges_exploitation',
        'Encours des crédits':        'encours_credits',
        'Crédits sains':              'credits_sains',
        'souffrance':                 'creances_souffrance',
        'irrécouvrables':             'irrecouvrables',
        'Taux de créances':           'taux_creances',
        'Taux de perte':              'taux_perte',
        'transformation':             'taux_transformation',
        'Résultat net':               'resultat_net',
        'Fonds propres':              'fonds_propres',
        'Immobilisations':            'immobilisations',
        'Actif total':                'actif_total',
        "agents":                     'nb_agents',
    }
    for row in t.rows:
        label = row.cells[0].text.strip()
        for key, field in label_map.items():
            if key.lower() in label.lower():
                valeurs = ind.get(field, {})
                if isinstance(valeurs, dict) and len(row.cells) >= 4:
                    _ecrire(row.cells[1], valeurs.get('p1', ''))
                    _ecrire(row.cells[2], valeurs.get('p2', ''))
                    _ecrire(row.cells[3], valeurs.get('p3', ''))
                break


# ══ TABLE 6 — Ressources (index 5) ══
def _table6(doc, mission):
    t = doc.tables[5]
    ressources = mission.get('ressources_credit', {})
    if not ressources: return
    lignes = ressources.get('lignes', [])
    LABELS = {
        'res-inst-fin':        'institutions financières',
        'res-cpt-ord-cred':    'ordinaires créditeurs',
        'res-dat':             'terme',
        'res-dep-gar':         'garantie reçus',
        'res-emprunts':        'Emprunts',
        'res-dep-membres':     'membres',
        'res-cpt-ord-membres': 'ordinaires créditeurs',
        'res-dat-recus':       'DAT reçus',
        'res-epargne-spec':    'régime spécial',
        'res-dep-gar2':        'garantie',
        'res-autres':          'Autres',
    }
    for ligne in lignes:
        label_key = LABELS.get(ligne.get('id', ''), '')
        if not label_key: continue
        for row in t.rows:
            if label_key.lower() in row.cells[0].text.lower() and len(row.cells) >= 4:
                _ecrire(row.cells[1], ligne.get('p1', ''))
                _ecrire(row.cells[2], ligne.get('p2', ''))
                _ecrire(row.cells[3], ligne.get('p3', ''))
                break


# ══ TABLE 10 — Suivi recommandations (index 9) ══
# ══ TABLE 10 — Suivi recommandations (index 9) ══
def _table10(doc, mission):
    t = doc.tables[9]
    suivis = mission.get('suivi_recommandations', [])
    data_rows = [r for r in t.rows if r.cells[0].text.strip().isdigit()]
    for i, s in enumerate(suivis):
        if i >= len(data_rows): break
        cells = data_rows[i].cells
        if len(cells) >= 4:
            # FIX : l'API /recommandations/ renvoie "description",
            # pas "recommandation" — d'où le fallback ci-dessous.
            recommandation = s.get('recommandation') or s.get('description', '')

            observation = s.get('observation')
            if not observation:
                # FIX : "observation" n'existe pas non plus dans le modèle
                # Recommandation — on reconstitue quelque chose de lisible
                # à partir de responsable + délai, qui existent réellement.
                parties = []
                if s.get('responsable'): parties.append(f"Resp. : {s['responsable']}")
                if s.get('delai'):       parties.append(f"Délai : {s['delai']}")
                observation = ' — '.join(parties)

            _ecrire(cells[1], recommandation)
            _ecrire(cells[2], s.get('statut', ''))
            _ecrire(cells[3], observation)
    t = doc.tables[9]
    suivis = mission.get('suivi_recommandations', [])
    data_rows = [r for r in t.rows if r.cells[0].text.strip().isdigit()]
    for i, s in enumerate(suivis):
        if i >= len(data_rows): break
        cells = data_rows[i].cells
        if len(cells) >= 4:
            _ecrire(cells[1], s.get('recommandation', ''))
            _ecrire(cells[2], s.get('statut', ''))
            _ecrire(cells[3], s.get('observation', ''))


# ══ TABLE 11 — Membres CA détaillé (index 10) ══
def _table11(doc, mission):
    t = doc.tables[10]
    membres = mission.get('organes', {}).get('ca', [])
    data_rows = [r for r in t.rows if r.cells[0].text.strip().isdigit()]
    for i, m in enumerate(membres):
        if i >= len(data_rows): break
        cells = data_rows[i].cells
        if len(cells) >= 7:
            _ecrire(cells[1], m.get('nom', ''))
            _ecrire(cells[2], m.get('fonction', ''))
            _ecrire(cells[3], m.get('debut', ''))
            _ecrire(cells[4], m.get('fin', ''))
            _ecrire(cells[5], m.get('profession', ''))
            _ecrire(cells[6], m.get('contact', ''))


# ══ TABLE 12 — Membres CC (index 11) ══
def _table12(doc, mission):
    t = doc.tables[11]
    membres = mission.get('organes', {}).get('cc', [])
    data_rows = [r for r in t.rows if r.cells[0].text.strip().isdigit()]
    for i, m in enumerate(membres):
        if i >= len(data_rows): break
        cells = data_rows[i].cells
        if len(cells) >= 7:
            _ecrire(cells[1], m.get('nom', ''))
            _ecrire(cells[2], m.get('fonction', ''))
            _ecrire(cells[3], m.get('debut', ''))
            _ecrire(cells[4], m.get('fin', ''))
            _ecrire(cells[5], m.get('profession', ''))
            _ecrire(cells[6], m.get('contact', ''))


# ══ TABLE 13 — Membres CS (index 12) ══
def _table13(doc, mission):
    t = doc.tables[12]
    membres = mission.get('organes', {}).get('cs', [])
    data_rows = [r for r in t.rows if r.cells[0].text.strip().isdigit()]
    for i, m in enumerate(membres):
        if i >= len(data_rows): break
        cells = data_rows[i].cells
        if len(cells) >= 7:
            _ecrire(cells[1], m.get('nom', ''))
            _ecrire(cells[2], m.get('fonction', ''))
            _ecrire(cells[3], m.get('debut', ''))
            _ecrire(cells[4], m.get('fin', ''))
            _ecrire(cells[5], m.get('profession', ''))
            _ecrire(cells[6], m.get('contact', ''))


# ══ TABLE 14 — Réunions organes (index 13) ══
def _table14(doc, mission):
    t = doc.tables[13]
    reunions = mission.get('reunions', {})
    if not reunions: return
    annees = reunions.get('annees', ['', '', ''])
    lignes = reunions.get('lignes', [])
    # Structure : L0=header, L1,L2,L3=années, L4=TOTAL
    # Colonnes : 0=Années, 1=AG, 2=CA, 3=CC, 4=CS
    for row_idx, annee in enumerate(annees[:3]):
        if row_idx+1 < len(t.rows) and annee:
            _ecrire(t.rows[row_idx+1].cells[0], annee)

    for col_idx in range(min(4, len(lignes))):
        ligne = lignes[col_idx]
        for row_idx, pk in enumerate(['a1', 'a2', 'a3']):
            if row_idx+1 < len(t.rows):
                row = t.rows[row_idx+1]
                if col_idx+1 < len(row.cells):
                    _ecrire(row.cells[col_idx+1], ligne.get(pk, ''))
        # Total
        if len(t.rows) > 4:
            total_row = t.rows[4]
            if col_idx+1 < len(total_row.cells):
                _ecrire(total_row.cells[col_idx+1], ligne.get('total', ''))


# ══ TABLE 15 — Gros épargnants (index 14) ══
def _table15(doc, mission):
    t = doc.tables[14]
    epargnants = mission.get('gros_epargnants', [])
    data_rows = [r for r in t.rows if r.cells[0].text.strip().isdigit()]
    for i, ep in enumerate(epargnants[:10]):
        if i >= len(data_rows): break
        cells = data_rows[i].cells
        if len(cells) >= 4:
            _ecrire(cells[1], ep.get('nom', ''))
            _ecrire(cells[2], ep.get('montant', ''))
            _ecrire(cells[3], ep.get('taux', ''))


# ══ TABLE 16 — Ressources détaillées (index 15) ══
def _table16(doc, mission):
    t = doc.tables[15]
    ressources = mission.get('ressources_credit', {})
    if not ressources: return
    lignes = ressources.get('lignes', [])
    LABELS = {
        'res-inst-fin':        'institutions financières',
        'res-cpt-ord-cred':    'Comptes ordinaires créditeurs',
        'res-dat':             'terme',
        'res-dep-gar':         'garantie reçus',
        'res-emprunts':        'Emprunts',
        'res-dep-membres':     'membres',
        'res-cpt-ord-membres': 'Comptes ordinaires',
        'res-dat-recus':       'DAT',
        'res-epargne-spec':    'spécial',
        'res-dep-gar2':        'garantie',
        'res-autres':          'Autres',
    }
    for ligne in lignes:
        label_key = LABELS.get(ligne.get('id', ''), '')
        if not label_key: continue
        for row in t.rows:
            if label_key.lower() in row.cells[0].text.lower() and len(row.cells) >= 4:
                _ecrire(row.cells[1], ligne.get('p1', ''))
                _ecrire(row.cells[2], ligne.get('p2', ''))
                _ecrire(row.cells[3], ligne.get('p3', ''))
                break


# ══ TABLE 17 — Production de prêts (index 16) ══
def _table17(doc, mission):
    t = doc.tables[16]
    prod = mission.get('production_prets', {})
    if not prod: return
    label_map = {
        'Nombre de prêts':     'prod-nb-prets',
        'mois':                'prod-nb-mois',
        'Production de prêts': 'prod-montant',
        'moyen':               'prod-moy-montant',
        'Évolution':           'prod-evol',
        'Variation':           'prod-var',
    }
    for row in t.rows:
        label = row.cells[0].text.strip()
        for key, fid in label_map.items():
            if key.lower() in label.lower():
                data = prod.get(fid, {})
                if isinstance(data, dict) and len(row.cells) >= 4:
                    _ecrire(row.cells[1], data.get('p1', ''))
                    _ecrire(row.cells[2], data.get('p2', ''))
                    _ecrire(row.cells[3], data.get('p3', ''))
                break


# ══ TABLE 18 — Portefeuille PAR (index 17) ══
def _table18(doc, mission):
    t = doc.tables[17]
    port = mission.get('portefeuille', {})
    if not port: return
    lignes = port.get('lignes', [])
    LABELS = {
        'port-brut':       'bruts',
        'port-sains':      'sains',
        'port-0j':         '0 jour',
        'port-1-30j':      '1 à 30',
        'port-31-60j':     '31 à 60',
        'port-61-90j':     '61 à 90',
        'port-souffrance': 'souffrance',
        'port-3-6m':       '3 à 6',
        'port-6-12m':      '6 à 12',
        'port-12-24m':     '12 à 24',
    }
    for ligne in lignes:
        label_key = LABELS.get(ligne.get('id', ''), '')
        if not label_key: continue
        for row in t.rows:
            cell_txt = row.cells[1].text if len(row.cells) > 1 else row.cells[0].text
            if label_key.lower() in cell_txt.lower() and len(row.cells) >= 5:
                _ecrire(row.cells[2], ligne.get('qte', ''))
                _ecrire(row.cells[4], ligne.get('val', ''))
                break


# ══ TABLE 19 — Prêts dirigeants (index 18) ══
def _table19(doc, mission):
    t = doc.tables[18]
    dirigeants = mission.get('prets_dirigeants', [])
    data_rows = [r for r in t.rows if r.cells[0].text.strip().isdigit()]
    for i, d in enumerate(dirigeants):
        if i >= len(data_rows): break
        cells = data_rows[i].cells
        if len(cells) >= 8:
            _ecrire(cells[1], d.get('titulaire', ''))
            _ecrire(cells[2], d.get('fonction', ''))
            _ecrire(cells[3], d.get('initial', ''))
            _ecrire(cells[4], d.get('restant', ''))
            _ecrire(cells[5], d.get('epargne', ''))
            _ecrire(cells[6], d.get('risque', ''))
            _ecrire(cells[7], d.get('retard', ''))


# ══ TABLE 20 — Synthèse par organe (index 19) ══
def _table20(doc, mission):
    t = doc.tables[19]
    synth = mission.get('synth_organes', [])
    data_rows = [r for r in t.rows
                 if r.cells[0].text.strip() and
                 r.cells[0].text.strip() not in ['Emprunteurs', 'Total']]
    for i, s in enumerate(synth):
        if i >= len(data_rows): break
        cells = data_rows[i].cells
        if len(cells) >= 7:
            _ecrire(cells[0], s.get('emprunteur', ''))
            _ecrire(cells[1], s.get('organe', ''))
            _ecrire(cells[2], s.get('initial', ''))
            _ecrire(cells[3], s.get('restant', ''))
            _ecrire(cells[4], s.get('epargne', ''))
            _ecrire(cells[5], s.get('risque', ''))
            _ecrire(cells[6], s.get('retard', ''))


# ══ TABLE 21 — Prêts personnel (index 20) ══
def _table21(doc, mission):
    t = doc.tables[20]
    personnel = mission.get('prets_personnel', [])
    data_rows = [r for r in t.rows if r.cells[0].text.strip().isdigit()]
    for i, p in enumerate(personnel):
        if i >= len(data_rows): break
        cells = data_rows[i].cells
        if len(cells) >= 8:
            _ecrire(cells[1], p.get('nom', ''))
            _ecrire(cells[2], p.get('fonction', ''))
            _ecrire(cells[3], p.get('initial', ''))
            _ecrire(cells[4], p.get('restant', ''))
            _ecrire(cells[5], p.get('epargne', ''))
            _ecrire(cells[6], p.get('risque', ''))
            _ecrire(cells[7], p.get('retard', ''))


# ══ TABLE 22 — Gros risques (index 21) ══
def _table22(doc, mission):
    t = doc.tables[21]
    gros_risques = mission.get('gros_risques', [])
    data_rows = [r for r in t.rows
                 if r.cells[0].text.strip() and
                 r.cells[0].text.strip() not in ['Emprunteurs', 'Total', 'TOTAL']]
    for i, gr in enumerate(gros_risques):
        if i >= len(data_rows): break
        cells = data_rows[i].cells
        if len(cells) >= 6:
            _ecrire(cells[0], gr.get('emprunteur', ''))
            _ecrire(cells[1], gr.get('initial', ''))
            _ecrire(cells[2], gr.get('restant', ''))
            _ecrire(cells[3], gr.get('epargne', ''))
            _ecrire(cells[4], gr.get('risque', ''))
            _ecrire(cells[5], gr.get('retard', ''))


# ══ TABLE 23 — Crédits virés en perte (index 22) ══
def _table23(doc, mission):
    t = doc.tables[22]
    perte = mission.get('credits_perte', {})
    if not perte: return
    LABELS = {
        'Nombre':       'perte-nb',
        'Montant':      'perte-mnt',
        'Récupération': 'perte-recup',
        'Encours':      'perte-enc',
        'Taux de perte':'perte-taux',
        'recouvrement': 'perte-recouvr',
    }
    for row in t.rows:
        label = row.cells[0].text.strip()
        for key, fid in LABELS.items():
            if key.lower() in label.lower():
                data = perte.get(fid, {})
                if isinstance(data, dict) and len(row.cells) >= 4:
                    _ecrire(row.cells[1], data.get('p1', ''))
                    _ecrire(row.cells[2], data.get('p2', ''))
                    _ecrire(row.cells[3], data.get('p3', ''))
                break


# ══ TABLE 24 — Prêts salariés TEG (index 23) ══
def _table24(doc, mission):
    t = doc.tables[23]
    salaries = mission.get('prets_salaries', [])
    data_rows = [r for r in t.rows
                 if r.cells[0].text.strip() and
                 r.cells[0].text.strip() not in ['Emprunteurs']]
    for i, s in enumerate(salaries):
        if i >= len(data_rows): break
        cells = data_rows[i].cells
        if len(cells) >= 5:
            _ecrire(cells[0], s.get('emprunteur', ''))
            _ecrire(cells[1], s.get('montant', ''))
            _ecrire(cells[2], s.get('taux', ''))
            _ecrire(cells[3], s.get('echeances', ''))
            _ecrire(cells[4], s.get('teg', ''))


# ══ TABLE 25 — Ratios prudentiels (index 24) ══
def _table25(doc, mission):
    t = doc.tables[24]
    ratios_data = mission.get('ratios', {})
    if not ratios_data: return
    periodes = ratios_data.get('periodes', ['', '', ''])
    lignes   = ratios_data.get('lignes', [])

    # Ligne 0 : colonnes 3,4,5 = périodes
    if len(t.rows) > 0:
        header = t.rows[0]
        for j, p in enumerate(periodes[:3]):
            if j+3 < len(header.cells):
                _ecrire(header.cells[j+3], p)

    # Lignes 1 à 10
    for ligne in lignes:
        num = int(ligne.get('num', 0))
        if 1 <= num < len(t.rows):
            row = t.rows[num]
            if len(row.cells) >= 7:
                _ecrire(row.cells[3], ligne.get('p1', ''))
                _ecrire(row.cells[4], ligne.get('p2', ''))
                _ecrire(row.cells[5], ligne.get('p3', ''))
                _ecrire(row.cells[6], ligne.get('obs', ''))


# ══ TABLE 26 — Lacunes & Recommandations (index 25) ══
SECTION_MAPPING = {
    "GOUVERNANCE":                          ["gov"],
    "CONTRÔLE INTERNE ET PLAN D'AFFAIRES":  ["ci"],
    "DISPOSITIF LBC/FT/FP":                 ["autres", "lbcft"],
    "GESTION DES RESSOURCES HUMAINES":      ["autres"],
    "GESTION DE L'ÉPARGNE":                 ["epg"],
    "GESTION DU CRÉDIT":                    ["cred"],
    "GESTION COMPTABLE ET FINANCIÈRE":      ["fin"],
    "SYSTÈME INFORMATIQUE ET SÉCURITÉ":     ["si", "sec"],
}

def _table26(doc, volets_par_code):
    t = doc.tables[25]
    sections = {}
    section_courante = None
    lignes_vides = []

    for i, row in enumerate(t.rows):
        texte = row.cells[0].text.strip().upper()
        # Ligne fusionnée = titre de section
        est_titre = (
            texte and len(texte) > 3 and
            all(c.text.strip().upper() == texte for c in row.cells) and
            texte not in ['NO', 'N°', 'LACUNES', 'RECOMMANDATIONS']
        )
        if est_titre:
            if section_courante:
                sections[section_courante] = lignes_vides
            section_courante = texte
            lignes_vides = []
        elif section_courante and not row.cells[0].text.strip():
            lignes_vides.append(i)

    if section_courante:
        sections[section_courante] = lignes_vides

    for titre, codes in SECTION_MAPPING.items():
        # Trouver la clé dans les sections
        section_key = None
        for key in sections:
            if titre.upper() in key or key in titre.upper():
                section_key = key
                break
        if not section_key:
            continue

        lignes_dispo = sections[section_key]
        constats = []
        for code in codes:
            if code in volets_par_code:
                data = volets_par_code[code].get('data', {})
                for cr in data.get('constats_recommandations', []):
                    if cr.get('constat') or cr.get('recommandation'):
                        constats.append(cr)

        for i, cr in enumerate(constats):
            if i >= len(lignes_dispo): break
            row = t.rows[lignes_dispo[i]]
            _ecrire(row.cells[0], str(i+1))
            _ecrire(row.cells[1], cr.get('constat', ''))
            _ecrire(row.cells[2], cr.get('recommandation', ''))
            if len(row.cells) > 3:
                _ecrire(row.cells[3], cr.get('echeance', ''))


# ══ FONCTION PRINCIPALE ══════════════════════════
def generer_rapport_global_bytes(mission_data: dict, volets_data: list, template_path: str) -> bytes:
    doc = Document(template_path)

    sfd          = mission_data.get('sfd', 'SFD INCONNU')
    reference    = mission_data.get('reference', 'N/A')
    date_mission = mission_data.get('date_mission', 'N/A')
    chef         = mission_data.get('chef_mission', 'Non renseigné')
    inspecteurs  = mission_data.get('inspecteurs', '')
    if isinstance(inspecteurs, list):
        inspecteurs = ', '.join(inspecteurs)
    reviseur     = mission_data.get('reviseur', 'Non désigné')

    # Remplacements textuels
    _remplacer_dans_doc(doc, {
        "indiquer la date de mission de contrôle":                 date_mission,
        "préciser le nom de la structure":                         sfd,
        "préciser les noms du chef de mission et des inspecteurs": f"{chef} (Chef de mission), {inspecteurs}",
        "(préciser le nom et la qualité)":                         reviseur,
        "(nom de la structure)":                                   sfd,
        "nom de la structure":                                     sfd,
        "DGTCP-DSFD-PR7-RAP-":                                    reference,
        "Date de la précédente mission, le cas échéant":          "Non applicable",
        "(Préciser les noms des inspecteurs)":                     inspecteurs,
    })

    # Organiser volets par code
    volets_par_code = {}
    for v in volets_data:
        code = v.get('volet_code')
        if code:
            volets_par_code[code] = v

    # ══ FIX — FUSION GÉNÉRIQUE DE TOUS LES VOLETS ══════════════════
    # Avant ce correctif, seules les données du volet crédit ('cred')
    # et épargne ('epg') étaient extraites et fusionnées dans
    # mission_data. Toute autre donnée saisie dans un volet (par ex.
    # personnel, indicateurs_financiers, suivi_recommandations, ou
    # toute autre clé future) n'atteignait donc jamais les tableaux
    # du rapport, même si elle avait bien été sauvegardée en base
    # dans donnees_volets.
    #
    # On fusionne ici, de façon générique, toutes les clés de premier
    # niveau du champ "data" de chaque volet validé dans mission_data
    # — sans écraser une valeur déjà présente et non vide sur la
    # mission (les colonnes infos_sfd/organes/reunions/ratios de la
    # table missions restent prioritaires si elles sont renseignées).
    for code, v in volets_par_code.items():
        data = v.get('data', {}) or {}
        for key, value in data.items():
            if value in (None, {}, [], ''):
                continue
            existant = mission_data.get(key)
            if existant in (None, {}, [], ''):
                mission_data[key] = value

    # Extraire données du volet crédit (règles spécifiques : certaines
    # clés comme gros_epargnants peuvent venir soit du volet crédit,
    # soit du volet épargne)
    cred = volets_par_code.get('cred', {}).get('data', {})
    epg  = volets_par_code.get('epg',  {}).get('data', {})

    mission_data['gros_epargnants']   = cred.get('gros_epargnants',  []) or epg.get('gros_epargnants', []) or mission_data.get('gros_epargnants', [])
    mission_data['ressources_credit'] = cred.get('ressources',       {}) or mission_data.get('ressources_credit', {})
    mission_data['production_prets']  = cred.get('production_prets', {}) or mission_data.get('production_prets', {})
    mission_data['portefeuille']      = cred.get('portefeuille',     {}) or mission_data.get('portefeuille', {})
    mission_data['prets_dirigeants']  = cred.get('prets_dirigeants', []) or mission_data.get('prets_dirigeants', [])
    mission_data['synth_organes']     = cred.get('synth_organes',    []) or mission_data.get('synth_organes', [])
    mission_data['prets_personnel']   = cred.get('prets_personnel',  []) or mission_data.get('prets_personnel', [])
    mission_data['gros_risques']      = cred.get('gros_risques',     []) or mission_data.get('gros_risques', [])
    mission_data['credits_perte']     = cred.get('credits_perte',    {}) or mission_data.get('credits_perte', {})
    mission_data['prets_salaries']    = cred.get('prets_salaries',   []) or mission_data.get('prets_salaries', [])

    # Remplir tous les tableaux
    try: _table2(doc, mission_data)
    except Exception as e: print(f"Table 2 erreur: {e}")

    try: _table3(doc, mission_data)
    except Exception as e: print(f"Table 3 erreur: {e}")

    try: _table4(doc, mission_data)
    except Exception as e: print(f"Table 4 erreur: {e}")

    try: _table5(doc, mission_data)
    except Exception as e: print(f"Table 5 erreur: {e}")

    try: _table6(doc, mission_data)
    except Exception as e: print(f"Table 6 erreur: {e}")

    try: _table10(doc, mission_data)
    except Exception as e: print(f"Table 10 erreur: {e}")

    try: _table11(doc, mission_data)
    except Exception as e: print(f"Table 11 erreur: {e}")

    try: _table12(doc, mission_data)
    except Exception as e: print(f"Table 12 erreur: {e}")

    try: _table13(doc, mission_data)
    except Exception as e: print(f"Table 13 erreur: {e}")

    try: _table14(doc, mission_data)
    except Exception as e: print(f"Table 14 erreur: {e}")

    try: _table15(doc, mission_data)
    except Exception as e: print(f"Table 15 erreur: {e}")

    try: _table16(doc, mission_data)
    except Exception as e: print(f"Table 16 erreur: {e}")

    try: _table17(doc, mission_data)
    except Exception as e: print(f"Table 17 erreur: {e}")

    try: _table18(doc, mission_data)
    except Exception as e: print(f"Table 18 erreur: {e}")

    try: _table19(doc, mission_data)
    except Exception as e: print(f"Table 19 erreur: {e}")

    try: _table20(doc, mission_data)
    except Exception as e: print(f"Table 20 erreur: {e}")

    try: _table21(doc, mission_data)
    except Exception as e: print(f"Table 21 erreur: {e}")

    try: _table22(doc, mission_data)
    except Exception as e: print(f"Table 22 erreur: {e}")

    try: _table23(doc, mission_data)
    except Exception as e: print(f"Table 23 erreur: {e}")

    try: _table24(doc, mission_data)
    except Exception as e: print(f"Table 24 erreur: {e}")

    try: _table25(doc, mission_data)
    except Exception as e: print(f"Table 25 erreur: {e}")

    try: _table26(doc, volets_par_code)
    except Exception as e: print(f"Table 26 erreur: {e}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()